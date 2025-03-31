from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
import mysql.connector
import os
from typing import Optional, List, Dict, Any
from dotenv import load_dotenv
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from fastapi.responses import StreamingResponse

# Import functions from sql_query_generator
from sql_query_generator import load_schema, create_prompt, execute_query

app = FastAPI(title="SQL Query Generator API")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Next.js default port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

load_dotenv()  # Load environment variables

# Configure Gemini API with debug print
api_key = os.getenv('GEMINI_API_KEY')
print(f"Using API key: {api_key}")  # Debug print
genai.configure(api_key=api_key)

# Database Configuration
DB_CONFIG = {
    'host': os.getenv('MYSQL_HOST', 'localhost'),
    'user': os.getenv('MYSQL_USER', 'root'),
    'password': os.getenv('MYSQL_PASSWORD', 'root'),
    'database': os.getenv('MYSQL_DATABASE', 'cybersecurity_incidents_3nf')
}

# Load schema once at startup
SCHEMA = load_schema()

class QueryRequest(BaseModel):
    query: str

class QueryResponse(BaseModel):
    sql_query: str
    columns: List[str]
    rows: List[List[Any]]

# Add this class for login request validation
class LoginRequest(BaseModel):
    email: str
    password: str

@app.post("/api/generate-query", response_model=QueryResponse)
async def generate_query(request: QueryRequest):
    try:
        # Generate SQL query
        model = genai.GenerativeModel('gemini-1.5-pro-001')
        prompt = create_prompt(SCHEMA, request.query)
        response = model.generate_content(prompt)
        
        # Clean response
        sql_query = response.text
        if '```sql' in sql_query:
            sql_query = sql_query.split('```sql')[1].split('```')[0].strip()
        else:
            sql_query = sql_query.strip()
            
        print(f"Generated SQL Query: {sql_query}")  # Debug print
            
        try:
            # Execute query using the DB_CONFIG
            conn = mysql.connector.connect(**DB_CONFIG)
            cursor = conn.cursor()
            cursor.execute(sql_query)
            
            # Fetch results
            columns = [desc[0] for desc in cursor.description]
            rows = cursor.fetchall()
            
            cursor.close()
            conn.close()
            
            if not columns or not rows:
                raise HTTPException(status_code=404, detail="No results found")
                
            return QueryResponse(
                sql_query=sql_query,
                columns=columns,
                rows=[[str(cell) for cell in row] for row in rows]  # Convert all cells to strings
            )
            
        except mysql.connector.Error as db_err:
            print(f"Database Error: {str(db_err)}")  # Debug print
            raise HTTPException(status_code=500, detail=f"Database error: {str(db_err)}")
            
    except Exception as e:
        print(f"General Error: {str(e)}")  # Debug print
        import traceback
        print(traceback.format_exc())  # Print full traceback
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate-insert-query")
async def generate_insert_query(request: dict):
    try:
        # Simplified prompt
        prompt = f"""
        Generate a MySQL INSERT query for incidents table with these values:
        Title: {request['title']}
        Type ID: {request['type']}
        Severity ID: {request['severity']}
        System ID: {request['affected_system']}
        Description: {request['description']}

        Use this format exactly:
        INSERT INTO incidents (incident_title, type_id, severity_id, system_id, reporter_id, status_id, discovered_at, reported_at, description) 
        VALUES ('title', type_id, severity_id, system_id, 1, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, 'description');
        """
        
        model = genai.GenerativeModel('gemini-1.5-pro-001')
        
        generation_config = {
            "temperature": 0.1,
            "top_p": 0.8,
            "top_k": 40,
            "max_output_tokens": 200
        }
        
        response = model.generate_content(
            prompt,
            generation_config=generation_config
        )
        
        if not response.text:
            # Fallback to manual query construction if AI fails
            sql_query = f"""
            INSERT INTO incidents 
            (incident_title, type_id, severity_id, system_id, reporter_id, status_id, discovered_at, reported_at, description)
            VALUES 
            ('{request['title']}', {request['type']}, {request['severity']}, {request['affected_system']}, 1, 1, 
            CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, '{request['description']}')
            """
            return {"sql_query": sql_query}
            
        sql_query = response.text.strip()
        
        # Clean up the query if it's wrapped in markdown
        if '```sql' in sql_query:
            sql_query = sql_query.split('```sql')[1].split('```')[0].strip()
            
        return {"sql_query": sql_query}
        
    except Exception as e:
        print(f"Query generation error: {str(e)}")  # Debug print
        # Fallback to manual query construction
        sql_query = f"""
        INSERT INTO incidents 
        (incident_title, type_id, severity_id, system_id, reporter_id, status_id, discovered_at, reported_at, description)
        VALUES 
        ('{request['title']}', {request['type']}, {request['severity']}, {request['affected_system']}, 1, 1, 
        CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, '{request['description']}')
        """
        return {"sql_query": sql_query}

@app.post("/api/execute-query")
async def execute_query(request: dict):
    try:
        query = request['query']
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        cursor.execute(query)
        incident_id = cursor.lastrowid  # Get the ID of the inserted row
        conn.commit()
        
        result = {
            "success": True, 
            "message": "Query executed successfully",
            "incident_id": incident_id
        }
        
        cursor.close()
        conn.close()
        
        return result
        
    except Exception as e:
        print(f"Error executing query: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/incident-types")
async def get_incident_types():
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT type_id, type_name FROM incident_types")
        types = cursor.fetchall()
        cursor.close()
        conn.close()
        return types
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/severity-levels")
async def get_severity_levels():
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT severity_id, severity_name FROM severity_levels")
        levels = cursor.fetchall()
        cursor.close()
        conn.close()
        return levels
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/affected-systems")
async def get_affected_systems():
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT system_id, system_name FROM affected_systems")
        systems = cursor.fetchall()
        cursor.close()
        conn.close()
        return systems
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/departments")
async def get_departments():
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT department_id, department_name FROM departments")
        departments = cursor.fetchall()
        cursor.close()
        conn.close()
        return departments
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/threat-guidelines/{incident_type_id}")
async def get_threat_guidelines(incident_type_id: int):
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # Join with threat_types to get guidelines for the incident type
        query = """
        SELECT g.guideline_category, g.guideline_description, g.risk_level
        FROM threat_guidelines g
        JOIN threat_types t ON g.threat_type_id = t.type_id
        JOIN incident_types i ON t.threat_name = i.type_name
        WHERE i.type_id = %s
        ORDER BY g.risk_level DESC, g.guideline_category
        """
        
        cursor.execute(query, (incident_type_id,))
        guidelines = cursor.fetchall()
        
        # Organize guidelines by category
        organized_guidelines = {
            'DO': [],
            'DONT': []
        }
        
        for guideline in guidelines:
            category = guideline['guideline_category']
            organized_guidelines[category].append({
                'description': guideline['guideline_description'],
                'risk_level': guideline['risk_level']
            })
        
        cursor.close()
        conn.close()
        
        return organized_guidelines
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/generate-report/{incident_id}")
async def generate_report(incident_id: int, request: Request):
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # Get complete incident details
        query = """
        SELECT 
            i.incident_id,
            i.incident_title,
            i.description,
            i.financial_impact,
            i.discovered_at,
            i.reported_at,
            r.first_name,
            r.last_name,
            st.severity_name,
            it.type_name,
            sys.system_name,
            d.department_name,
            s.status_name
        FROM incidents i
        LEFT JOIN reporters r ON i.reporter_id = r.reporter_id
        LEFT JOIN severity_levels st ON i.severity_id = st.severity_id
        LEFT JOIN incident_types it ON i.type_id = it.type_id
        LEFT JOIN affected_systems sys ON i.system_id = sys.system_id
        LEFT JOIN departments d ON r.department_id = d.department_id
        LEFT JOIN incident_statuses s ON i.status_id = s.status_id
        WHERE i.incident_id = %s
        """
        
        cursor.execute(query, (incident_id,))
        incident = cursor.fetchone()
        
        if not incident:
            raise HTTPException(status_code=404, detail="Incident not found")
            
        # Create report document
        doc = Document()
        
        # Add title
        doc.add_heading('Cybersecurity Incident Report', 0)
        
        # Add incident number in a box/rectangle
        incident_number = doc.add_heading(f'Incident Report #{incident_id}', level=1)
        incident_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add Incident Summary section
        doc.add_heading('Incident Summary', level=1)
        
        # Create summary table
        table = doc.add_table(rows=10, cols=2)
        table.style = 'Table Grid'
        
        # Fill table with incident details
        rows = [
            ('Incident Title', incident['incident_title'] or 'Untitled Incident'),
            ('Incident Type', incident['type_name']),
            ('Severity Level', incident['severity_name']),
            ('Status', incident['status_name']),
            ('Affected System', incident['system_name']),
            ('Department', incident['department_name']),
            ('Discovered At', incident['discovered_at'].strftime('%Y-%m-%d %H:%M:%S')),
            ('Reported At', incident['reported_at'].strftime('%Y-%m-%d %H:%M:%S')),
            ('Reporter', f"{incident['first_name']} {incident['last_name']}"),
            ('Financial Impact', f"${incident['financial_impact']:,.2f}" if incident['financial_impact'] else 'N/A')
        ]
        
        for i, (key, value) in enumerate(rows):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = str(value)
            
        # Add Incident Description
        doc.add_heading('Incident Description', level=1)
        doc.add_paragraph(incident['description'] or 'No description provided')
        
        # Get and add Response Guidelines
        cursor.execute("""
            SELECT g.guideline_category, g.guideline_description, g.risk_level
            FROM threat_guidelines g
            JOIN threat_types t ON g.threat_type_id = t.type_id
            JOIN incident_types i ON t.threat_name = i.type_name
            WHERE i.type_id = (
                SELECT type_id FROM incidents WHERE incident_id = %s
            )
            ORDER BY g.risk_level DESC, g.guideline_category
        """, (incident_id,))
        
        guidelines = cursor.fetchall()
        
        # Add Response Guidelines section
        doc.add_heading('Response Guidelines', level=1)
        
        # Add Do's
        doc.add_heading("Do's:", level=2)
        for guideline in guidelines:
            if guideline['guideline_category'] == 'DO':
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(f"{guideline['guideline_description']} ")
                p.add_run(f"(Risk Level: {guideline['risk_level']})").italic = True
                
        # Add Don'ts
        doc.add_heading("Don'ts:", level=2)
        for guideline in guidelines:
            if guideline['guideline_category'] == 'DONT':
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(f"{guideline['guideline_description']} ")
                p.add_run(f"(Risk Level: {guideline['risk_level']})").italic = True
        
        # Save to memory stream
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                'Content-Disposition': f'attachment; filename=incident_report_{incident_id}.docx'
            }
        )
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error generating report: {str(e)}"
        )
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()

@app.put("/api/incidents/{incident_id}/assign")
async def assign_incident(incident_id: int, request: Request):
    try:
        data = await request.json()
        member_id = data.get('member_id')
        status = data.get('status', 'IN_PROGRESS')
        
        if not member_id:
            raise HTTPException(
                status_code=400,
                detail="member_id is required"
            )
            
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        try:
            # First get the status_id
            cursor.execute(
                "SELECT status_id FROM incident_statuses WHERE status_name = %s",
                (status,)
            )
            status_result = cursor.fetchone()
            
            if not status_result:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid status: {status}"
                )
                
            status_id = status_result[0]
            
            # Update the incident with the assigned team member
            query = """
            UPDATE incidents 
            SET assigned_to = %s, 
                status_id = %s,
                updated_at = CURRENT_TIMESTAMP
            WHERE incident_id = %s
            """
            
            cursor.execute(query, (member_id, status_id, incident_id))
            
            if cursor.rowcount == 0:
                raise HTTPException(
                    status_code=404,
                    detail=f"Incident {incident_id} not found"
                )
                
            conn.commit()
            
            return {
                "success": True,
                "message": "Security team member assigned successfully",
                "incident_id": incident_id,
                "member_id": member_id,
                "status_id": status_id
            }
            
        except mysql.connector.Error as db_error:
            conn.rollback()
            raise HTTPException(
                status_code=500,
                detail=f"Database error: {str(db_error)}"
            )
            
    except HTTPException as http_error:
        raise http_error
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Unexpected error: {str(e)}"
        )
    finally:
        if 'cursor' in locals():
            cursor.close()
        if 'conn' in locals():
            conn.close()

@app.get("/api/security-team")
async def get_security_team():
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        # Simplified query first to debug
        query = """
        SELECT 
            st.member_id,
            st.first_name,
            st.last_name,
            st.email,
            sr.role_name
        FROM security_team st
        JOIN security_roles sr ON st.role_id = sr.role_id
        """
        
        cursor.execute(query)
        team_members = cursor.fetchall()
        
        cursor.close()
        conn.close()
        
        return team_members
        
    except Exception as e:
        print(f"Security team error: {str(e)}")  # Debug print
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/update-password/{reporter_id}")
async def update_password(reporter_id: int, password: str):
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor()
        
        query = """
        UPDATE reporters 
        SET password = %s 
        WHERE reporter_id = %s
        """
        
        cursor.execute(query, (password, reporter_id))
        conn.commit()
        
        cursor.close()
        conn.close()
        
        return {"message": "Password updated successfully"}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/login")
async def login(request: LoginRequest):
    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        
        query = """
        SELECT reporter_id, first_name, last_name, email, department_id 
        FROM reporters 
        WHERE email = %s AND password = %s
        """
        
        cursor.execute(query, (request.email, request.password))
        user = cursor.fetchone()
        
        cursor.close()
        conn.close()
        
        if user:
            return {
                "success": True,
                "user": user
            }
        else:
            raise HTTPException(status_code=401, detail="Invalid credentials")
            
    except Exception as e:
        print(f"Login error: {str(e)}")  # Debug print
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 